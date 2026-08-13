#!/usr/bin/env python3
"""Build the editable, double-anonymized journal submission package."""

from __future__ import annotations

import csv
import html
import shutil
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
OUT = PAPER / "submission"
FIGURES = OUT / "figures"


def rows(name: str) -> list[dict[str, str]]:
    with (ROOT / "tables" / name).open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def esc(value: object) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ")


def md_table(headers: list[str], body: list[list[object]]) -> str:
    lines = [
        "| " + " | ".join(map(esc, headers)) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    lines.extend("| " + " | ".join(map(esc, row)) + " |" for row in body)
    return "\n".join(lines)


def p(value: str) -> str:
    number = float(value)
    if number < 0.001:
        return "< .001"
    return f"{number:.3f}".lstrip("0")


def main_tables() -> str:
    t1 = rows("table1_characteristics.csv")
    t1_headers = list(t1[0])
    t1_md = md_table(t1_headers, [[row[h] for h in t1_headers] for row in t1])

    t2a = rows("table2a_missing_rates.csv")
    t2a_md = md_table(
        ["Variable", "Present, n", "Missing, n (%)"],
        [
            [
                row["variable"],
                f'{int(row["n_present"]):,}',
                f'{int(row["n_missing"]):,} ({float(row["pct_missing"]):.1f}%)',
            ]
            for row in t2a
        ],
    )
    t2b = rows("table2b_missing_by_year.csv")
    t2b_md = md_table(
        ["Year", "N", "Gender", "Age", "Education", "Employment", "Prior record", "Severity"],
        [
            [
                row["year"],
                f'{int(row["n"]):,}',
                *[
                    f'{float(row[column]):.1f}%'
                    for column in (
                        "gender_pct_missing",
                        "age_pct_missing",
                        "education_level_pct_missing",
                        "employed_pct_missing",
                        "prior_criminal_pct_missing",
                        "severity_pct_missing",
                    )
                ],
            ]
            for row in t2b
        ],
    )

    t3 = rows("table3_bivariate.csv")
    t3_md = md_table(
        ["Variable", "Comparison", "Test", "Statistic", "p", "Effect size"],
        [
            [
                row["variable"],
                " vs. ".join(filter(None, [row["group1"], row["group2"]])),
                row["test"],
                f'{float(row["statistic"]):.2f}',
                p(row["p_value"]),
                f'{float(row["effect_size"]):.3f} ({row["effect_type"]})',
            ]
            for row in t3
        ],
    )

    t4 = rows("table4_model_hierarchy.csv")
    order = ["female", "age", "age_sq", "education_level", "employed", "prior_criminal"]
    labels = {
        "female": "Female",
        "age": "Age",
        "age_sq": "Age squared",
        "education_level": "Education level",
        "employed": "Employed",
        "prior_criminal": "Prior criminal record",
    }
    by_model = {(row["model"], row["variable"]): row for row in t4}
    models = ["Model 1 - Total Association", "Model 2 - Same-Crime"]
    t4_body: list[list[str]] = []
    for variable in order:
        line = [labels[variable]]
        for model in models:
            row = by_model[(model, variable)]
            line.extend(
                [
                    f'{float(row["coefficient"]):.2f} ({float(row["se"]):.2f})',
                    p(row["p"]),
                ]
            )
        t4_body.append(line)
    t4_md = md_table(
        ["Variable", "Model 1 B (SE)", "p", "Model 2 B (SE)", "p"], t4_body
    )

    t5 = [row for row in rows("table5_robustness.csv") if row["variable"] == "employed"]
    t5_md = md_table(
        ["Specification", "B", "SE", "p", "N", "Outcome/note"],
        [
            [
                row["model"],
                f'{float(row["coefficient"]):.3f}',
                f'{float(row["se"]):.3f}',
                p(row["p"]),
                f'{int(row["n"]):,}',
                row["notes"],
            ]
            for row in t5
        ],
    )

    return f"""
## Tables

**Table 1. Sample characteristics.**

{t1_md}

*Note.* Sample A is the preregistered complete-case adult sample. Sample B omits the age requirement and excludes defendants known to be juveniles.

**Table 2. Missing data in the cleaned sample.**

*Panel A. Overall missingness.*

{t2a_md}

*Panel B. Percentage missing by decision year.*

{t2b_md}

**Table 3. Bivariate associations with sentence severity.**

{t3_md}

**Table 4. Validated core regression models (N = 29,132).**

{t4_md}

*Note.* OLS coefficients with court-clustered standard errors in parentheses. Both models include year and court fixed effects. Model 1 includes broad crime category; Model 2 includes 91 crime-article groups. Model 1 R-squared = .080; Model 2 R-squared = .625.

**Table 5. Employment coefficient across robustness specifications.**

{t5_md}

*Note.* Cross-sentence models use month-equivalent severity unless indicated. Fine-only and imprisonment-only coefficients use logged sanction amounts and are interpreted proportionally.
"""


def supplement() -> str:
    parts = ["# Supplementary Material", ""]
    for title, headers, body in supplement_tables():
        parts.extend([f"## {title}", "", md_table(headers, body), ""])
    return "\n".join(parts)


def supplement_tables() -> list[tuple[str, list[str], list[list[str]]]]:
    predictor = {
        "female": "Female",
        "age": "Age",
        "age_sq": "Age squared",
        "education_level": "Education level",
        "employed": "Employed",
        "prior_criminal": "Prior criminal record",
    }

    holm = rows("table4c_holm_bonferroni.csv")
    a1 = (
        "Table A1. Holm-Bonferroni correction for Model 1",
        ["Hypothesis", "Predictor", "p", "Holm threshold", "Decision"],
        [[row["hypothesis"], predictor.get(row["variable"], row["variable"]),
          p(row["p"]), f'{float(row["threshold"]):.4f}',
          "Reject" if row["reject"] == "True" else "Do not reject"] for row in holm],
    )

    variance = rows("table4b_variance_decomposition.csv")
    a2 = (
        "Table A2. Variance decomposition",
        ["Block", "R-squared", "Delta R-squared"],
        [[row["block"], f'{float(row["r_squared"]):.3f}',
          f'{float(row["delta_r_squared"]):.3f}'] for row in variance],
    )

    stage1 = rows("table5b_two_part_stage1.csv")
    a3 = (
        "Table A3. Two-part imprisonment model, stage 1",
        ["Predictor", "Logit B", "SE", "z", "p", "Odds ratio", "95% OR CI"],
        [[predictor.get(row["variable"], row["variable"]),
          f'{float(row["coefficient"]):.3f}', f'{float(row["se"]):.3f}',
          f'{float(row["z"]):.2f}', p(row["p"]), f'{float(row["odds_ratio"]):.2f}',
          f'{float(row["or_ci_low"]):.2f} to {float(row["or_ci_high"]):.2f}']
         for row in stage1],
    )

    robustness = rows("table5_robustness.csv")
    a4 = (
        "Table A4. Full robustness estimates",
        ["Specification", "Predictor", "B", "SE", "p", "95% CI", "R-squared", "N", "Note"],
        [[row["model"], predictor.get(row["variable"], row["variable"]),
          f'{float(row["coefficient"]):.3f}', f'{float(row["se"]):.3f}', p(row["p"]),
          f'{float(row["ci_low"]):.3f} to {float(row["ci_high"]):.3f}',
          f'{float(row["r_squared"]):.3f}' if row["r_squared"] else "—",
          f'{int(row["n"]):,}', row["notes"]] for row in robustness],
    )

    registry = rows("sample_registry.csv")
    a5 = (
        "Table A5. Sample registry",
        ["Sample", "Description", "N", "Required fields"],
        [[row["sample_key"], row["description"], f'{int(row["n"]):,}',
          row["required_columns"] or "None"] for row in registry],
    )

    metrics = rows("validation_metrics.csv")
    a6 = (
        "Table A6. Validation metrics",
        ["Field", "Metric", "Value", "N", "Threshold", "Gate result"],
        [[row["metric_group"].capitalize(), row["metric"].replace("_", " "),
          f'{float(row["value"]):.3f}', f'{int(row["n"]):,}',
          f'{float(row["threshold"]):.3f}' if row["threshold"] else "—",
          {"True": "Pass", "False": "Fail"}.get(row["gate_pass"], "Not applicable")]
         for row in metrics],
    )
    return [a1, a2, a3, a4, a5, a6]


def build_supplement_pdf(path: Path) -> None:
    page_width, _ = landscape(A4)
    margin = 12 * mm
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "SupplementTitle", parent=styles["Title"], fontName="Times-Bold",
        fontSize=16, leading=19, alignment=TA_CENTER,
    )
    heading_style = ParagraphStyle(
        "SupplementHeading", parent=styles["Heading2"], fontName="Times-Bold",
        fontSize=10, leading=12, spaceAfter=6,
    )
    cell_style = ParagraphStyle(
        "SupplementCell", parent=styles["BodyText"], fontName="Times-Roman",
        fontSize=5.5, leading=6.5,
    )
    header_style = ParagraphStyle(
        "SupplementHeader", parent=cell_style, fontName="Times-Bold",
    )
    doc = SimpleDocTemplate(
        str(path), pagesize=landscape(A4), leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin, title="Anonymous Supplementary Material",
        author="",
    )
    story = [
        Paragraph("Supplementary Material", title_style),
        Paragraph("Demographic Disparities in Criminal Sentencing: Evidence from Mongolian Courts", styles["Heading2"]),
        Paragraph("Asian Journal of Criminology — anonymous peer-review version", styles["BodyText"]),
        Spacer(1, 8),
    ]
    available_width = page_width - 2 * margin
    tables = supplement_tables()
    for index, (title, headers, body) in enumerate(tables):
        if index:
            story.append(PageBreak())
        story.append(Paragraph(title, heading_style))
        wrapped = [
            [Paragraph(html.escape(str(value)), header_style) for value in headers],
            *[[Paragraph(html.escape(str(value)), cell_style) for value in row] for row in body],
        ]
        weights = [
            2.5 if header in {"Note", "Required fields"}
            else 1.8 if header in {"Specification", "Description"}
            else 1.25 if header in {"Predictor", "Hypothesis", "95% CI", "95% OR CI"}
            else 0.8
            for header in headers
        ]
        unit = available_width / sum(weights)
        table = Table(wrapped, colWidths=[weight * unit for weight in weights], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.append(table)
    doc.build(story)


def embed_figures(path: Path) -> None:
    doc = Document(path)
    image_names = {
        1: "fig1_demographics.png",
        2: "fig2_sentences.png",
        3: "fig3_bivariate.png",
        4: "fig4_crime_composition.png",
        5: "fig5_model_hierarchy.png",
        6: "fig6_employment_robustness.png",
    }
    captions = {
        number: next(
            paragraph for paragraph in doc.paragraphs
            if paragraph.text.strip().startswith(f"Fig. {number}")
        )
        for number in range(1, 7)
    }
    for number, caption in captions.items():
        image_paragraph = caption.insert_paragraph_before()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = image_paragraph.add_run()
        shape = run.add_picture(
            str(ROOT / "figures" / image_names[number]), width=Inches(6.25)
        )
        shape._inline.docPr.set("descr", caption.text)
    doc.save(path)


def add_page_number(paragraph) -> None:
    paragraph.alignment = 2
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def format_docx(path: Path, *, anonymous: bool, double_spaced: bool = True) -> None:
    doc = Document(path)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
        add_page_number(section.footer.paragraphs[0])
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing_rule = (
        WD_LINE_SPACING.DOUBLE if double_spaced else WD_LINE_SPACING.SINGLE
    )
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for table in doc.tables:
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
    doc.core_properties.author = "" if anonymous else "Robert Ritz"
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.save(path)


def convert(markdown: str, output: Path, *, citeproc: bool = False) -> None:
    args = ["--standalone"]
    if citeproc:
        args.extend(
            [
                "--citeproc",
                f"--bibliography={ROOT / 'references.bib'}",
                "--csl=https://www.zotero.org/styles/apa",
            ]
        )
    pypandoc.convert_text(
        markdown,
        "docx",
        format="markdown",
        outputfile=str(output),
        extra_args=args,
    )


def build() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    FIGURES.mkdir(parents=True, exist_ok=True)

    manuscript = (PAPER / "manuscript.md").read_text(encoding="utf-8")
    tables = main_tables()
    manuscript = manuscript.replace("## Figure Captions", tables + "\n\n## Figure Captions", 1)
    manuscript_docx = OUT / "manuscript_anonymized.docx"
    convert(manuscript, manuscript_docx, citeproc=True)
    format_docx(manuscript_docx, anonymous=True)
    embed_figures(manuscript_docx)

    title_page = """# Demographic Disparities in Criminal Sentencing: Evidence from Mongolian Courts

Robert Ritz  
American University of Mongolia, Ulaanbaatar, Mongolia  
Corresponding author: Robert Ritz, me@robertritz.com  

## Statements and Declarations

**Funding.** The author received no specific grant for this work. The research was conducted as part of the author's institutional role at the American University of Mongolia.

**Competing interests.** The author has no relevant financial or non-financial interests to disclose.

**Ethics approval.** No formal ethics committee approval was obtained. The study is a secondary analysis of judicial decisions made publicly available by Mongolia's judiciary. No individuals were recruited or contacted, no intervention occurred, and the manuscript reports only aggregate results. Direct identifiers have been removed from the released analysis dataset.

**Consent to participate and consent for publication.** Not applicable. The study uses public judicial records and reports no identifiable individual-level findings.

**Pre-registration.** The study was preregistered on the Open Science Framework on February 3, 2026. The registration is publicly available at https://osf.io/sjvek.

**Data and code availability.** The analysis code, aggregate validation evidence, and a deidentified replication dataset are publicly available at https://github.com/robertritz/mongolia-sentencing-disparities. Source decisions remain available from shuukh.mn, subject to the judiciary's continued availability and terms.

**Author contributions.** Robert Ritz: conceptualization, methodology, software, data curation, formal analysis, investigation, validation, visualization, writing—original draft, writing—review and editing, and project administration.

**Use of generative AI.** xAI's `grok-4-1-fast-non-reasoning` model (Grok 4.1 Fast) was used for structured extraction from Mongolian judicial decisions. Three separate prediction-blinded Codex sessions recoded disjoint portions of the fixed 300-decision validation sample under a protocol locked before model predictions were exposed. Codex also assisted with code generation, analysis checking, and editorial revision. The author reviewed the source-coding protocol, code, outputs, and manuscript, made the final research decisions, and remains accountable for the work.
"""
    title_docx = OUT / "title_page.docx"
    convert(title_page, title_docx)
    format_docx(title_docx, anonymous=False, double_spaced=False)

    cover_letter = """# Cover Letter

Dear Editor,

Please consider the manuscript “Demographic Disparities in Criminal Sentencing: Evidence from Mongolian Courts” for publication in the *Asian Journal of Criminology*.

The study analyzes 80,827 published criminal first-instance decisions from Mongolia and reports validated adult complete-case models covering 29,132 cases. It contributes a large-scale quantitative account of sentencing patterns in Mongolia, tests preregistered demographic hypotheses, and shows that employment status is strongly associated with sentence routing and severity even within criminal-code article. A prediction-blinded 300-decision source audit validated employment extraction while identifying and removing an invalid probation category and unvalidated aggravating and mitigating factor measures.

The manuscript fits the journal’s focus on empirical criminology in Asian contexts and contributes both substantive evidence from an understudied legal system and a transparent measurement lesson for computational court-record research.

An earlier public-facing article (https://robertritz.com/2026/03/12/are-mongolian-courts-biased-criminal.html) reported preliminary analyses from this project. The submitted manuscript is a substantially revised scholarly analysis based on a locked publication validation audit, corrected sample restrictions, exclusion of an invalid sentence category, and removal of the unvalidated factor-dependent results. The public article has been corrected to identify the earlier results as preliminary and to report only findings supported by the publication audit.

This manuscript is original and is not under consideration elsewhere.

Thank you for your consideration.

Sincerely,  
Robert Ritz  
American University of Mongolia  
me@robertritz.com
"""
    cover_docx = OUT / "cover_letter.docx"
    convert(cover_letter, cover_docx)
    format_docx(cover_docx, anonymous=False, double_spaced=False)

    supplement_docx = OUT / "supplementary_material.docx"
    convert(supplement(), supplement_docx)
    format_docx(supplement_docx, anonymous=True)
    build_supplement_pdf(OUT / "ESM_1.pdf")

    for number in range(1, 7):
        source = ROOT / "figures" / f"Fig{number}.eps"
        shutil.copy2(source, FIGURES / source.name)

    print(f"Built submission package at {OUT}")


if __name__ == "__main__":
    build()
