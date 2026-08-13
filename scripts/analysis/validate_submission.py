#!/usr/bin/env python3
"""Fail fast on publication-critical evidence and manuscript inconsistencies."""

from __future__ import annotations

import csv
import hashlib
import json
import re
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[2]


def check(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)
    print(f"PASS: {message}")


def main() -> None:
    validation = ROOT / "data" / "validation"
    lock = json.loads((validation / "gold_labels.lock.json").read_text())
    coding_sheet = validation / "manual_coding_sheet.csv"
    if coding_sheet.exists():
        digest = hashlib.sha256(coding_sheet.read_bytes()).hexdigest()
        check(digest == lock["coding_sheet_sha256"], "gold-label hash matches the lock record")
    else:
        check(
            re.fullmatch(r"[0-9a-f]{64}", lock["coding_sheet_sha256"]) is not None,
            "public package contains the gold-label lock digest",
        )
    check(lock["row_count"] == 300, "gold-label lock contains 300 cases")

    metrics = list(csv.DictReader((ROOT / "tables" / "validation_metrics.csv").open()))
    by_metric = {(row["metric_group"], row["metric"]): row for row in metrics}
    check(float(by_metric[("employment", "accuracy")]["value"]) >= 0.90, "employment accuracy passes")
    check(float(by_metric[("employment", "kappa")]["value"]) >= 0.80, "employment kappa passes")
    check(by_metric[("aggravating", "present_f1")]["gate_pass"] == "False", "failed aggravating measure remains flagged")
    check(by_metric[("mitigating", "present_f1")]["gate_pass"] == "False", "failed mitigating measure remains flagged")

    data = pd.read_parquet(ROOT / "data" / "processed" / "sentencing_clean.parquet")
    check(len(data) == 69_527, "canonical cleaned data has 69,527 rows")
    check("probation" not in set(data["sentence_type"].dropna()), "contaminated probation label is absent")

    samples = {row["sample_key"]: int(row["n"]) for row in csv.DictReader((ROOT / "tables" / "sample_registry.csv").open())}
    check(samples == {"full": 69527, "sample_a": 29132, "sample_b": 35995, "two_part_stage1": 33600, "two_part_stage2": 5521}, "sample registry matches canonical counts")

    bivariate = {row["variable"]: row for row in csv.DictReader((ROOT / "tables" / "table3_bivariate.csv").open())}
    check(round(abs(float(bivariate["Employed (H4)"]["effect_size"])), 3) == 0.394, "bivariate employment effect uses adult Sample A")
    check(round(float(bivariate["Age non-linear (H3)"]["statistic"]), 2) == 0.28, "bivariate age test uses adult Sample A")

    manuscript = (ROOT / "paper" / "manuscript.md").read_text(encoding="utf-8")
    abstract = re.search(r"## Abstract\n\n(.+?)\n\n\*\*Keywords", manuscript, re.S).group(1)
    abstract_words = re.findall(r"\b[\w’-]+\b", abstract)
    check(150 <= len(abstract_words) <= 250, f"abstract length is {len(abstract_words)} words")
    keyword_text = re.search(r"\*\*Keywords:\*\* (.+)", manuscript).group(1)
    keywords = [item.strip() for item in keyword_text.split(",")]
    check(4 <= len(keywords) <= 6, f"keyword count is {len(keywords)}")

    stale = ["29,847", "29,259", "75,323", "36,829", "93 groups", "Model 3 (Pre-Registered)", "publicly available for replication", "d = 0.393", "F = 0.89"]
    check(not [token for token in stale if token in manuscript], "manuscript contains no superseded sample/model claims")
    check("prediction-blinded manual" not in manuscript.lower(), "Codex recoding is not described as manual validation")
    check(
        "Robert Ritz" not in manuscript
        and "me@robertritz.com" not in manuscript
        and "github.com/robertritz" not in manuscript
        and "osf.io/sjvek" not in manuscript,
        "manuscript is author-anonymized",
    )

    bib = (ROOT / "references.bib").read_text(encoding="utf-8")
    bib_keys = set(re.findall(r"@\w+\{([^,]+),", bib))
    cited = set(re.findall(r"@([\w:-]+)", manuscript))
    check(cited <= bib_keys, f"all {len(cited)} citation keys exist in the bibliography")
    check("fledgling2014courts" not in cited, "misattributed Mongolia chapter key is absent")

    for number in range(1, 6):
        check(re.search(rf"Table {number}\b", manuscript) is not None, f"Table {number} is cited")
    for number in range(1, 7):
        check(re.search(rf"Figure {number}\b", manuscript) is not None, f"Figure {number} is cited")
        eps_path = ROOT / "figures" / f"Fig{number}.eps"
        check(eps_path.exists(), f"Figure {number} EPS exists")
        eps_text = eps_path.read_text(encoding="latin-1")
        bbox = re.search(
            r"^%%HiResBoundingBox:\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)",
            eps_text,
            re.M,
        )
        check(bbox is not None, f"Figure {number} has an EPS bounding box")
        x0, y0, x1, y1 = map(float, bbox.groups())
        check(x1 - x0 <= 494 and y1 - y0 <= 664, f"Figure {number} fits journal dimensions")
        check("/FontType 42 def" in eps_text, f"Figure {number} embeds TrueType fonts")
        check(
            re.search(rf"\*\*Fig\. {number}\*\* [^\n]+[^.\n]$", manuscript, re.M)
            is not None,
            f"Figure {number} caption follows journal convention",
        )
    check("Online Resource 1" in manuscript, "supplement is cited as Online Resource 1")

    models = list(csv.DictReader((ROOT / "tables" / "table4_model_hierarchy.csv").open()))
    m1_emp = next(row for row in models if row["model"].startswith("Model 1") and row["variable"] == "employed")
    m2_emp = next(row for row in models if row["model"].startswith("Model 2") and row["variable"] == "employed")
    check(round(float(m1_emp["coefficient"]), 2) == -6.28, "Model 1 employment coefficient is canonical")
    check(round(float(m2_emp["coefficient"]), 2) == -2.83, "Model 2 employment coefficient is canonical")
    robustness = list(csv.DictReader((ROOT / "tables" / "table5_robustness.csv").open()))
    fine_tail = next(row for row in robustness if row["model"] == "Exclude fines above personal statutory maximum" and row["variable"] == "employed")
    check(int(fine_tail["n"]) == 29_124 and round(float(fine_tail["coefficient"]), 2) == -6.23, "statutory fine-tail sensitivity is canonical")

    out = ROOT / "paper" / "submission"
    expected = [
        "manuscript_anonymized.docx",
        "title_page.docx",
        "cover_letter.docx",
        "supplementary_material.docx",
        "ESM_1.pdf",
    ]
    check(all((out / name).exists() for name in expected), "all submission documents exist")
    title_doc = Document(out / "title_page.docx")
    title_text = "\n".join(paragraph.text for paragraph in title_doc.paragraphs)
    check(
        "Bayaraa Zorigt" in title_text
        and "Vice Rector" in title_text
        and "formal ethics review was not required" in title_text,
        "independent institutional ethics determination is stated on the title page",
    )
    manuscript_doc = Document(out / "manuscript_anonymized.docx")
    doc_text = "\n".join(paragraph.text for paragraph in manuscript_doc.paragraphs)
    check("@starr" not in doc_text and "References" in doc_text, "DOCX citations are rendered and references included")
    check("Robert Ritz" not in doc_text and "me@robertritz.com" not in doc_text, "DOCX manuscript is anonymized")
    check(len(manuscript_doc.inline_shapes) == 6, "all six figures are embedded in the manuscript")
    descriptions = [shape._inline.docPr.get("descr", "") for shape in manuscript_doc.inline_shapes]
    check(
        all(description.startswith(f"Fig. {number}") for number, description in enumerate(descriptions, 1)),
        "embedded figures have descriptive alternative text",
    )
    check(len(Document(out / "supplementary_material.docx").tables) == 6, "supplement contains all six supporting tables")
    supplement_pdf = (out / "ESM_1.pdf").read_bytes()
    check(
        supplement_pdf.startswith(b"%PDF") and len(supplement_pdf) > 10_000,
        "journal-format supplementary PDF exists",
    )
    check(
        b"Robert Ritz" not in supplement_pdf and b"me@robertritz.com" not in supplement_pdf,
        "supplementary PDF is anonymized",
    )
    check(all((out / "figures" / f"Fig{number}.eps").exists() for number in range(1, 7)), "all six submission figures exist")

    print("\nSubmission preflight passed.")


if __name__ == "__main__":
    main()
